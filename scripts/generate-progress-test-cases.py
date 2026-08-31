"""
Generate CampusCare Web System Software Progress Test Cases DOCX
organized by role: Landing Page, Authentication, Admin, Nurse, Physician, Dentist.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt

# Use original blank template if available; fall back to current downloads file.
TEMPLATE_CANDIDATES = [
    Path(r"c:\Users\Anjob\Downloads\CAMPUSCARE WEB PROGRESS TEST CASES - TEMPLATE.docx"),
    Path(r"c:\Users\Anjob\Desktop\campuscare-group-project\docs\CAMPUSCARE WEB PROGRESS TEST CASES - TEMPLATE.docx"),
    Path(r"c:\Users\Anjob\Downloads\CAMPUSCARE WEB PROGRESS TEST CASES.docx"),
]
OUTPUT_DOWNLOADS = Path(
    r"c:\Users\Anjob\Downloads\CAMPUSCARE WEB PROGRESS TEST CASES.docx"
)
OUTPUT_PROJECT = Path(
    r"c:\Users\Anjob\Desktop\campuscare-group-project\docs\CAMPUSCARE WEB PROGRESS TEST CASES.docx"
)

TestCase = tuple[str, str, str, str]  # module, scenario, action, actual_input

SECTIONS: list[tuple[str, list[TestCase]]] = [
    (
        "LANDING PAGE",
        [
            (
                "Landing Page",
                "Verify the public landing page loads",
                "1. Open the CampusCare web application URL.\n2. Observe the landing page.",
                "Browser: Chrome/Edge; URL: /landing",
            ),
            (
                "Landing Page",
                "Verify landing page navigation and Login link",
                "1. On the landing page, locate the top navigation.\n2. Click Login.\n3. Confirm navigation to the login page.",
                "N/A",
            ),
            (
                "Landing Page",
                "Verify Learn More scrolls to About section",
                "1. On the landing page hero, click Learn More.\n2. Observe the page scroll position.",
                "N/A",
            ),
            (
                "Landing Page",
                "Verify About, Features, How It Works, and FAQ sections are visible",
                "1. On the landing page, scroll through About, Features, How It Works, and FAQ.\n2. Confirm each section displays content.",
                "N/A",
            ),
            (
                "Landing Page",
                "Verify footer legal links open",
                "1. Scroll to the landing page footer.\n2. Open Privacy Policy, Terms of Service, and Data Privacy links.",
                "N/A",
            ),
        ],
    ),
    (
        "AUTHENTICATION PAGE",
        [
            (
                "Authentication",
                "Verify login page displays required fields",
                "1. Navigate to /login.\n2. Observe the sign-in form.",
                "N/A",
            ),
            (
                "Authentication",
                "Verify empty email validation on login",
                "1. On /login, leave Work email empty.\n2. Click Send one-time password.",
                "Work email: (empty)",
            ),
            (
                "Authentication",
                "Verify invalid email format validation",
                "1. On /login, enter an invalid email.\n2. Click Send one-time password.",
                "Work email: not-an-email",
            ),
            (
                "Authentication",
                "Verify valid staff email proceeds to OTP step",
                "1. On /login, enter a valid active staff email.\n2. Click Send one-time password.",
                "Work email: active staff account email",
            ),
            (
                "Authentication",
                "Verify OTP step displays after email submission",
                "1. Complete the email step with a valid staff email.\n2. Observe the Verify your login screen.",
                "Valid staff email",
            ),
            (
                "Authentication",
                "Verify OTP resend cooldown",
                "1. On the OTP step, click Resend.\n2. Observe the 60-second cooldown timer.",
                "Valid staff email; 6-digit OTP screen",
            ),
            (
                "Authentication",
                "Verify invalid OTP shows error",
                "1. On the OTP step, enter an incorrect 6-digit code.\n2. Submit verification.",
                "OTP: 000000",
            ),
            (
                "Authentication",
                "Verify successful staff login redirects to role home",
                "1. Sign in with valid staff email and correct OTP.\n2. Observe redirect after login.",
                "Valid nurse, physician, dentist, or admin account",
            ),
            (
                "Authentication",
                "Verify unauthenticated access to staff page redirects to login",
                "1. While logged out, open /nurse/dashboard.\n2. Observe redirect to login.",
                "URL: /nurse/dashboard",
            ),
            (
                "Authentication",
                "Verify pending invite page for users without clinic membership",
                "1. Sign in with an account that has no clinic membership.\n2. Observe Invite pending page.",
                "Invited but not yet assigned account",
            ),
            (
                "Authentication",
                "Verify logout confirmation dialog",
                "1. Sign in as any staff user.\n2. Open the user menu.\n3. Click Log out.",
                "Any active staff account",
            ),
            (
                "Authentication",
                "Verify logout cancel keeps user signed in",
                "1. Open Log Out? dialog.\n2. Click Cancel.\n3. Verify the user remains signed in.",
                "Any active staff account",
            ),
            (
                "Authentication",
                "Verify logout confirm returns to login page",
                "1. Open Log Out? dialog.\n2. Click Log Out.\n3. Observe redirect to /login.",
                "Any active staff account",
            ),
            (
                "Authentication",
                "Verify display login entry from staff login logo",
                "1. On /login, triple-click the CampusCare logo within 800 ms.\n2. Observe navigation to /display-login.",
                "N/A",
            ),
            (
                "Authentication",
                "Verify queue display login page fields",
                "1. Open /display-login.\n2. Observe Email, Password, Open display, and Staff login link.",
                "N/A",
            ),
            (
                "Authentication",
                "Verify activated account notice on login page",
                "1. Open /login?activated=1.\n2. Observe the activated account notice message.",
                "URL: /login?activated=1",
            ),
            (
                "Authentication",
                "Verify queue display login opens public display board",
                "1. On /display-login, enter valid display account credentials.\n2. Click Open display.",
                "Display account email and password",
            ),
        ],
    ),
    (
        "ADMIN",
        [
            (
                "Admin Navigation",
                "Verify admin sidebar shows correct menu items",
                "1. Sign in as Admin.\n2. Review sidebar.\n3. Confirm Dashboard, Reports, Announcements, and Clinic Staff are shown.",
                "Admin account",
            ),
            (
                "Admin Navigation",
                "Verify admin cannot access nurse queue page directly",
                "1. Sign in as Admin.\n2. Open /nurse/queue manually.\n3. Observe redirect to admin home.",
                "Admin account",
            ),
            (
                "Admin Dashboard",
                "Verify admin dashboard loads after login",
                "1. Sign in as Admin.\n2. Open Dashboard from the sidebar.",
                "Admin account",
            ),
            (
                "Admin Dashboard",
                "Verify summary cards display clinic metrics",
                "1. On Admin Dashboard, review summary cards.\n2. Confirm consultations, patients served, pending requests, queue, medical consultations, and dental consultations.",
                "Admin account",
            ),
            (
                "Admin Dashboard",
                "Verify consultation overview chart period toggle",
                "1. On Admin Dashboard, locate Consultation overview.\n2. Switch Daily, Weekly, and Monthly.",
                "Admin account",
            ),
            (
                "Admin Dashboard",
                "Verify quick link Create Announcement navigates correctly",
                "1. On Admin Dashboard, click Create Announcement.",
                "Admin account",
            ),
            (
                "Admin Dashboard",
                "Verify quick link View Reports navigates correctly",
                "1. On Admin Dashboard, click View Reports.",
                "Admin account",
            ),
            (
                "Admin Dashboard",
                "Verify quick link Manage Staff navigates correctly",
                "1. On Admin Dashboard, click Manage Staff.",
                "Admin account",
            ),
            (
                "Admin Dashboard",
                "Verify Staff on duty panel shows role badges",
                "1. On Admin Dashboard, review Staff on duty.\n2. Confirm Active, Invited, and Inactive badges appear where applicable.",
                "Admin account",
            ),
            (
                "Admin Reports",
                "Verify admin reports page and tabs display",
                "1. Sign in as Admin.\n2. Open Reports.\n3. Review tabs: Clinic Operations, Consultations, Patients Served, Queue Performance, Consultation Requests, Documents Issued.",
                "Admin account",
            ),
            (
                "Admin Reports",
                "Verify admin date range and quick range filters",
                "1. On Admin Reports, set From/To dates.\n2. Click Last 7d, Last 30d, and Last 90d.\n3. Observe report updates.",
                "From: first day of month; To: today",
            ),
            (
                "Admin Reports",
                "Verify admin Print, PDF, CSV, and Excel export buttons",
                "1. On Admin Reports, confirm Print, PDF, CSV, and Export Excel are visible.\n2. Run one export action.",
                "Admin account",
            ),
            (
                "Admin Announcements",
                "Verify admin announcements management table loads",
                "1. Sign in as Admin.\n2. Open Announcements.\n3. Review Published, Scheduled, Drafts, and Total stat cards.",
                "Admin account",
            ),
            (
                "Admin Announcements",
                "Verify admin can search and filter announcements",
                "1. Search title, body, audience, author, or status.\n2. Change status filter among all, published, scheduled, draft.",
                "Search: Health",
            ),
            (
                "Admin Announcements",
                "Verify admin Add announcement opens form",
                "1. Click Add announcement.\n2. Review Title, Body, Audience, Status, and Scheduled at fields.",
                "Admin account",
            ),
            (
                "Admin Announcements",
                "Verify admin Publish confirmation and toast",
                "1. On a draft announcement, click Publish.\n2. Confirm Publish Announcement?.\n3. Observe Announcement Published toast.",
                "Draft announcement",
            ),
            (
                "Admin Announcements",
                "Verify admin can view published announcement feed",
                "1. Open Announcements.\n2. Scroll to published news feed.\n3. Confirm published cards display for admin.",
                "Admin account",
            ),
            (
                "Admin Announcements",
                "Verify admin Edit, Unpublish, Archive, and Delete actions",
                "1. Execute Edit on an announcement and save.\n2. Test Unpublish, Archive, and Delete with their confirmation dialogs and toasts.",
                "Existing announcement",
            ),
            (
                "Admin Clinic Staff",
                "Verify clinic staff page loads",
                "1. Sign in as Admin.\n2. Open Clinic Staff from sidebar.",
                "Admin account",
            ),
            (
                "Admin Clinic Staff",
                "Verify staff summary cards and search filters",
                "1. Review Staff accounts, Active, Invited, Inactive cards.\n2. Search users and filter by role and status.",
                "Search: staff name or email",
            ),
            (
                "Admin Clinic Staff",
                "Verify Invite staff with confirmation",
                "1. Click Invite staff.\n2. Complete Name, email, and role.\n3. Confirm Create Staff Account?.\n4. Observe Staff Account Created toast.",
                "Name: Test Nurse; Email: test.nurse@nu-dasma.edu.ph; Role: Nurse",
            ),
            (
                "Admin Clinic Staff",
                "Verify Import staff opens import sheet",
                "1. Click Import staff.\n2. Review Excel template columns: Full Name, Email, Employee ID, License No., Role.",
                "Admin account",
            ),
            (
                "Admin Clinic Staff",
                "Verify Edit, Enable, Disable, and Delete staff actions",
                "1. Edit a staff account and save.\n2. Test Enable, Disable, and Delete with confirmation dialogs and toasts.",
                "Existing staff account",
            ),
            (
                "Admin Clinic Staff",
                "Verify Resend invite for invited accounts",
                "1. Locate an Invited account.\n2. Click Resend invite.\n3. Observe success feedback.",
                "Invited staff account",
            ),
            (
                "Admin Clinic Staff",
                "Verify bulk staff selection actions",
                "1. Select multiple staff rows.\n2. Review bulk Activate, Deactivate, Archive, and Delete actions.\n3. Execute one action with confirmation.",
                "Two or more staff rows",
            ),
            (
                "Admin Settings",
                "Verify admin Profile and Settings opens from user menu",
                "1. Open user menu.\n2. Click Profile and Settings.",
                "Admin account",
            ),
            (
                "Admin Settings",
                "Verify admin consultation capacity settings",
                "1. On Settings, open Consultation Capacity.\n2. Update values and save.\n3. Observe Consultation Slots Updated toast.",
                "Admin account",
            ),
            (
                "Admin Settings",
                "Verify admin office hours settings",
                "1. On Settings, open Office Hours.\n2. Update clinic hours and save.\n3. Observe Clinic Schedule Updated toast.",
                "Admin account",
            ),
            (
                "Admin Settings",
                "Verify admin light and dark mode toggle",
                "1. Open user menu.\n2. Switch Light mode / Dark mode.\n3. Observe theme change.",
                "Admin account",
            ),
            (
                "Admin Session",
                "Verify admin idle session warning after 4 minutes",
                "1. Sign in as Admin.\n2. Remain inactive for 4 minutes.\n3. Observe Session Inactivity Detected dialog.",
                "Admin account; no activity",
            ),
            (
                "Admin Session",
                "Verify admin session lock and Continue Session unlock",
                "1. Allow session to lock after inactivity warning.\n2. Click Continue Session.\n3. Verify same page is restored.",
                "Admin account",
            ),
            (
                "Admin Negative",
                "Verify duplicate staff invite shows error feedback",
                "1. Invite staff using an email already in the directory.\n2. Observe error toast or validation message.",
                "Existing staff email",
            ),
        ],
    ),
    (
        "NURSE STAFF",
        [
            (
                "Nurse Navigation",
                "Verify nurse sidebar shows operations and clinical items",
                "1. Sign in as Nurse.\n2. Confirm Dashboard, Reports, Announcements, Consultation Requests, Queue Management, Patient Records, Consultations, and Medical Documents appear.",
                "Nurse account",
            ),
            (
                "Nurse Navigation",
                "Verify nurse is redirected away from physician routes",
                "1. Sign in as Nurse.\n2. Open /physician/dashboard manually.\n3. Observe redirect to nurse home.",
                "Nurse account",
            ),
            (
                "Nurse Dashboard",
                "Verify nurse dashboard loads",
                "1. Sign in as Nurse.\n2. Open Dashboard.",
                "Nurse account",
            ),
            (
                "Nurse Dashboard",
                "Verify Register walk-in opens walk-in sheet",
                "1. On Nurse Dashboard, click Register walk-in.",
                "Nurse account",
            ),
            (
                "Nurse Dashboard",
                "Verify Open queue navigates to nurse queue page",
                "1. On Nurse Dashboard, click Open queue.",
                "Nurse account",
            ),
            (
                "Nurse Dashboard",
                "Verify At a glance KPI cards display",
                "1. Review intake, pending, waiting, and served cards.",
                "Nurse account",
            ),
            (
                "Nurse Dashboard",
                "Verify Today's queue and Triage panels display",
                "1. Review Today's queue and Triage & follow-up panels.\n2. Review Recent activity feed.",
                "Nurse account",
            ),
            (
                "Nurse Consultation Requests",
                "Verify Consultation Requests page loads",
                "1. Open Consultation Requests from sidebar.\n2. Review Pending, Waitlisted, Rescheduled, and Declined stat cards.",
                "Nurse account",
            ),
            (
                "Nurse Consultation Requests",
                "Verify search by Student ID and status filter",
                "1. Search by Student ID.\n2. Filter by Pending, Waitlisted, Rescheduled, and All statuses.",
                "Student ID: 2023-171863",
            ),
            (
                "Nurse Consultation Requests",
                "Verify View opens request details",
                "1. Click View on a request.\n2. Review request details dialog.",
                "Any visible request",
            ),
            (
                "Nurse Consultation Requests",
                "Verify Approve pending request with confirmation",
                "1. Click Approve on a Pending request.\n2. Confirm Approve Consultation?.\n3. Observe Request Approved toast and updated status.",
                "Pending request",
            ),
            (
                "Nurse Consultation Requests",
                "Verify Approve cancel leaves request unchanged",
                "1. Click Approve.\n2. Cancel Approve Consultation? dialog.\n3. Confirm status remains Pending.",
                "Pending request",
            ),
            (
                "Nurse Consultation Requests",
                "Verify Admit waitlisted request",
                "1. Locate a Waitlisted request.\n2. Click Admit and complete confirmation.",
                "Waitlisted request",
            ),
            (
                "Nurse Consultation Requests",
                "Verify Reschedule request with confirmation",
                "1. Click Reschedule.\n2. Enter valid date/time.\n3. Confirm Reschedule Consultation?.\n4. Observe Request Rescheduled toast.",
                "Valid future date/time",
            ),
            (
                "Nurse Consultation Requests",
                "Verify Decline request with confirmation",
                "1. Click Decline.\n2. Confirm Decline Consultation Request?.\n3. Observe Request Declined toast.",
                "Pending request",
            ),
            (
                "Nurse Consultation Requests",
                "Verify empty state when no requests match filters",
                "1. Apply filters or search with no matching records.\n2. Observe No consultation requests message.",
                "Search: nonexistent ID",
            ),
            (
                "Nurse Queue Management",
                "Verify queue status filter options",
                "1. Open Status filter.\n2. Review All statuses, Waiting, Called, Completed, No Show, Expired.",
                "Nurse queue page",
            ),
            (
                "Nurse Queue Management",
                "Verify empty queue message",
                "1. Apply filters with no matching tickets.\n2. Observe No tickets match these filters message.",
                "Filtered empty queue",
            ),
            (
                "Nurse Queue Management",
                "Verify Rejoin end of queue when available",
                "1. Locate a row with Rejoin end of queue.\n2. Execute action.\n3. Verify patient returns to waiting state.",
                "Eligible skipped/no-show row",
            ),
            (
                "Nurse Queue Management",
                "Verify nurse queue page title and summary cards",
                "1. Open Queue Management.\n2. Confirm title Nurse queue and Waiting, Current, Completed today cards.",
                "Nurse account",
            ),
            (
                "Nurse Queue Management",
                "Verify Register walk-in from queue page",
                "1. Click Register walk-in.\n2. Complete required fields.\n3. Submit.",
                "Patient ID: 2023-171863; Reason: Headache",
            ),
            (
                "Nurse Queue Management",
                "Verify walk-in required fields validation",
                "1. Open Register walk-in.\n2. Leave required fields empty.\n3. Attempt submit.",
                "Empty required fields",
            ),
            (
                "Nurse Queue Management",
                "Verify Call next calls waiting patient",
                "1. With waiting patients in queue, click Call next.\n2. Observe Patient Called toast and status change.",
                "Queue with waiting patients",
            ),
            (
                "Nurse Queue Management",
                "Verify nurse lane switcher tabs",
                "1. Switch Needs intake, Specialty, and Exceptions lanes.",
                "Nurse account",
            ),
            (
                "Nurse Queue Management",
                "Verify search by ID Number",
                "1. Search queue using a known ID Number.",
                "ID Number: 2023-171863",
            ),
            (
                "Nurse Queue Management",
                "Verify Verify check-in action",
                "1. Open Queue actions.\n2. Select Verify check-in.\n3. Complete verification.",
                "Patient awaiting check-in",
            ),
            (
                "Nurse Queue Management",
                "Verify Intake and assign specialty action",
                "1. Open Queue actions.\n2. Select Intake & assign specialty.\n3. Complete intake sheet and save.",
                "Patient needing intake",
            ),
            (
                "Nurse Queue Management",
                "Verify intake records vital signs",
                "1. Enter blood pressure, heart rate, temperature, and related vitals in intake.\n2. Save.\n3. Observe Vital Signs Recorded toast if shown.",
                "BP: 120/80; HR: 72; Temp: 36.5",
            ),
            (
                "Nurse Queue Management",
                "Verify Assign queue number action",
                "1. Open Queue actions.\n2. Select Assign queue number.\n3. Complete prompt.",
                "Eligible queue row",
            ),
            (
                "Nurse Queue Management",
                "Verify Skip patient with confirmation",
                "1. Open Queue actions.\n2. Select Skip.\n3. Confirm Skip Patient?.\n4. Observe Patient Skipped toast.",
                "Waiting patient",
            ),
            (
                "Nurse Queue Management",
                "Verify Mark no-show with confirmation",
                "1. Select Mark no-show (need 2 calls).\n2. Confirm Remove Patient from Queue?.\n3. Observe Patient Removed toast.",
                "Patient who failed to respond",
            ),
            (
                "Nurse Queue Management",
                "Verify Transfer to Physician, Dentist, or Nurse",
                "1. Open Queue actions on transferable row.\n2. Transfer to another station.\n3. Confirm queue updates.",
                "Eligible queue row",
            ),
            (
                "Nurse Public Queue Display",
                "Verify public queue display shows station boards after nurse queue activity",
                "1. Sign in as Nurse in one browser.\n2. Open /queue-management/display in another browser or screen.\n3. From nurse queue, call a patient.\n4. Confirm the display board updates.",
                "Nurse account; display URL open",
            ),
            (
                "Nurse Public Queue Display",
                "Verify public queue display shows waiting counts",
                "1. With patients waiting in nurse queue, view /queue-management/display.\n2. Confirm station sections and total waiting count appear.",
                "Active clinic queue",
            ),
            (
                "Nurse Public Queue Display",
                "Verify recently served updates on public display after nurse completes patient",
                "1. Complete a patient from nurse queue.\n2. View /queue-management/display.\n3. Confirm patient appears in recently served.",
                "Nurse account; display open",
            ),
            (
                "Nurse Public Queue Display",
                "Verify public display loads without staff sidebar",
                "1. Open /queue-management/display.\n2. Confirm full-screen queue board without staff navigation shell.",
                "Display URL",
            ),
            (
                "Nurse Consultations",
                "Verify nurse consultations page loads",
                "1. Open Consultations.\n2. Review Open today, In nurse queue, Vitals in progress, and Completed today cards.",
                "Nurse account",
            ),
            (
                "Nurse Consultations",
                "Verify nurse consultation filters and Create consultation",
                "1. Use patient search, status filter, provider type filter, and date range filter.\n2. Click Create consultation.",
                "Nurse account",
            ),
            (
                "Nurse Consultations",
                "Verify View opens consultation summary dialog",
                "1. Click View on a consultation.\n2. Review summary content.",
                "Any listed consultation",
            ),
            (
                "Nurse Patient Records",
                "Verify patient records page loads with stat cards",
                "1. Open Patient Records.\n2. Review Patients on file, Visited this month, Flagged allergies, and Documents cards.",
                "Nurse account",
            ),
            (
                "Nurse Patient Records",
                "Verify patient type filter and search",
                "1. Filter by All types, Student, Faculty, and Employee.\n2. Search by patient name or ID.",
                "Search: 2023-171863",
            ),
            (
                "Nurse Patient Records",
                "Verify Import patients opens import sheet",
                "1. Click Import patients.\n2. Review import instructions and sample format.",
                "Nurse account",
            ),
            (
                "Nurse Patient Records",
                "Verify Profile, History, Documents, and Update medical actions",
                "1. On a patient row, open Profile, History, Documents, and Update medical.\n2. Review each sheet content.",
                "Patient with history and documents",
            ),
            (
                "Nurse Patient Records",
                "Verify edit patient and save shows toast",
                "1. Open Profile.\n2. Edit allowed fields.\n3. Save.\n4. Observe Patient Record Saved toast.",
                "Updated contact or demographic field",
            ),
            (
                "Nurse Patient Records",
                "Verify empty patient search result",
                "1. Search for a nonexistent patient.\n2. Observe No patients found message.",
                "Search: ZZZNONEXISTENT",
            ),
            (
                "Nurse Medical Documents",
                "Verify nurse medical documents view-only access",
                "1. Open Medical Documents.\n2. Confirm search and view are available.\n3. Confirm issue/generate actions are restricted.",
                "Nurse account",
            ),
            (
                "Nurse Reports",
                "Verify nurse reports page and filters",
                "1. Open Reports.\n2. Set date range, consultation type, patient type, personnel, and status filters.\n3. Confirm daily consultation, queue performance, and consultation request reports appear.",
                "Nurse account",
            ),
            (
                "Nurse Reports",
                "Verify nurse Print/PDF and CSV export; Excel hidden",
                "1. Confirm Print/PDF and Export CSV are visible.\n2. Confirm Export Excel is not shown.",
                "Nurse account",
            ),
            (
                "Nurse Announcements",
                "Verify nurse announcements management",
                "1. Open Announcements.\n2. Use Add, Edit, Publish, Unpublish, Archive, and Delete with confirmation dialogs.",
                "Nurse account",
            ),
            (
                "Nurse Duty Management",
                "Verify Start Duty confirmation and Available status",
                "1. Click Start Duty.\n2. Confirm Start Duty?.\n3. Observe Available badge and Duty Started toast.",
                "Nurse off duty",
            ),
            (
                "Nurse Duty Management",
                "Verify End Duty confirmation and Not Available status",
                "1. Click End Duty.\n2. Confirm End Duty?.\n3. Observe Not Available badge and Duty Ended toast.",
                "Nurse on duty",
            ),
            (
                "Nurse Break",
                "Verify Go on Break confirmation and overlay",
                "1. Click Break.\n2. Confirm Go on Break?.\n3. Observe On Break overlay and blocked interaction.",
                "Nurse with Available status",
            ),
            (
                "Nurse Break",
                "Verify Resume Work confirmation restores nurse workflow",
                "1. Click Resume Work.\n2. Confirm Resume Work?.\n3. Verify overlay closes and current page remains open.",
                "Nurse on break",
            ),
            (
                "Nurse Session",
                "Verify nurse idle session warning and lock",
                "1. Remain inactive for 4 minutes.\n2. Observe Session Inactivity Detected.\n3. Allow lock and unlock with Continue Session.",
                "Nurse account",
            ),
            (
                "Nurse Settings",
                "Verify nurse profile, notifications, schedule, and capacity settings",
                "1. Open Profile and Settings.\n2. Review personal info, notification toggles, staff schedule, and capacity settings.\n3. Save a change and observe success toast.",
                "Nurse account",
            ),
            (
                "Nurse Responsive",
                "Verify nurse mobile sidebar and queue usability",
                "1. Resize below 768px.\n2. Open sidebar and navigate.\n3. Open Queue Management and confirm controls remain usable.",
                "Mobile viewport; Nurse account",
            ),
        ],
    ),
    (
        "PHYSICIAN STAFF",
        [
            (
                "Physician Navigation",
                "Verify physician sidebar hides Consultation Requests",
                "1. Sign in as Physician.\n2. Review sidebar.\n3. Confirm Consultation Requests is not listed.",
                "Physician account",
            ),
            (
                "Physician Navigation",
                "Verify physician is redirected away from nurse requests page",
                "1. Sign in as Physician.\n2. Open /nurse/requests manually.\n3. Observe redirect to physician home.",
                "Physician account",
            ),
            (
                "Physician Dashboard",
                "Verify physician dashboard loads",
                "1. Sign in as Physician.\n2. Open Dashboard.",
                "Physician account",
            ),
            (
                "Physician Dashboard",
                "Verify Open queue navigates to physician station queue",
                "1. Click Open queue.\n2. Confirm title Your station queue.",
                "Physician account",
            ),
            (
                "Physician Dashboard",
                "Verify KPI cards and queue panels",
                "1. Review waiting, serving, completed, and appointments cards.\n2. Review Now serving, Waiting list with vitals strip, and Recently served panels.",
                "Physician account",
            ),
            (
                "Physician Queue Management",
                "Verify physician queue summary cards and search",
                "1. Open Queue Management.\n2. Review Waiting, Current, and Completed today cards.\n3. Search by ID Number.",
                "Physician account",
            ),
            (
                "Physician Queue Management",
                "Verify Register walk-in is hidden for physician",
                "1. Open Queue Management as Physician.\n2. Confirm Register walk-in is not shown.",
                "Physician account",
            ),
            (
                "Physician Queue Management",
                "Verify Call / recall action",
                "1. Open Queue actions.\n2. Select Call / recall.\n3. Observe Patient Called or Patient Recalled toast.",
                "Waiting or previously called patient",
            ),
            (
                "Physician Queue Management",
                "Verify Start consultation with confirmation",
                "1. Open Queue actions.\n2. Click Start consultation.\n3. Confirm Start Consultation?.\n4. Observe navigation to consultation page.",
                "Waiting patient in physician queue",
            ),
            (
                "Physician Queue Management",
                "Verify Complete, Skip, and Mark no-show actions",
                "1. Test Complete, Skip (with Skip Patient? confirmation), and Mark no-show (with Remove Patient from Queue? confirmation).",
                "Active or waiting queue row",
            ),
            (
                "Physician Public Queue Display",
                "Verify public display updates when physician calls patient",
                "1. Open /queue-management/display on a second screen.\n2. As Physician, call a patient from queue.\n3. Confirm display board updates.",
                "Physician account; display open",
            ),
            (
                "Physician Public Queue Display",
                "Verify public display shows recently served after physician completes patient",
                "1. Complete a patient from physician queue.\n2. View /queue-management/display.\n3. Confirm patient appears in recently served.",
                "Physician account; display open",
            ),
            (
                "Physician Public Queue Display",
                "Verify vitals appear on physician waiting list after nurse intake",
                "1. After nurse records vitals, open Physician Dashboard waiting list.\n2. Confirm vitals summary appears for patient.",
                "Patient with recorded vitals",
            ),
            (
                "Physician Consultations",
                "Verify physician consultations page loads",
                "1. Open Consultations.\n2. Review Queued today, Waiting to be called, In consultation, and Completed today cards.",
                "Physician account",
            ),
            (
                "Physician Consultations",
                "Verify physician consultation filters; Create consultation hidden",
                "1. Use search, status, and date range filters.\n2. Confirm Create consultation is not shown.",
                "Physician account",
            ),
            (
                "Physician Consultations",
                "Verify Open navigates to live consultation page",
                "1. Click Open on an active consultation.\n2. Observe /physician/consultation/[id] page.",
                "Ongoing or callable consultation",
            ),
            (
                "Physician Consultations",
                "Verify consultation form fields editable while ongoing",
                "1. Edit chief complaint, symptoms, assessment, diagnosis, treatment, and prescription.\n2. Click Save draft.\n3. Observe Consultation Saved toast.",
                "Sample text: Headache for 2 days",
            ),
            (
                "Physician Consultations",
                "Verify Complete consultation with confirmation",
                "1. Click Complete consultation.\n2. Confirm Complete consultation? dialog.\n3. Observe Consultation Completed toast.",
                "Ongoing consultation",
            ),
            (
                "Physician Consultations",
                "Verify completed consultation is read-only",
                "1. Re-open a completed consultation.\n2. Confirm fields are read-only and title shows Medical Record (Completed).",
                "Completed consultation",
            ),
            (
                "Physician Consultations",
                "Verify issue document from active consultation visit",
                "1. On live consultation page, issue a Medical Certificate.\n2. Complete wizard and finalize if prompted.",
                "Physician account; Medical Certificate",
            ),
            (
                "Physician Patient Records",
                "Verify physician patient records access",
                "1. Open Patient Records.\n2. Use search, filters, Profile, History, Documents, and Update medical actions.",
                "Physician account",
            ),
            (
                "Physician Consultations",
                "Verify Complete consultation cancel leaves consultation ongoing",
                "1. Click Complete consultation.\n2. Cancel Complete consultation? dialog.\n3. Confirm consultation remains ongoing.",
                "Ongoing consultation",
            ),
            (
                "Physician Medical Documents",
                "Verify document preview dialog matches issued layout",
                "1. Click Preview on an issued Medical Certificate.\n2. Review Document preview dialog content and layout.",
                "Issued medical certificate",
            ),
            (
                "Physician Medical Documents",
                "Verify print from preview opens portrait half-bond print dialog",
                "1. In Document preview, click Print.\n2. Observe browser print dialog.",
                "Issued medical certificate",
            ),
            (
                "Physician Medical Documents",
                "Verify physician can issue, preview, print, void, and delete documents",
                "1. Open Medical Documents.\n2. Issue Medical Certificate.\n3. Preview and Print.\n4. Test Void and Delete with confirmations.",
                "Physician account",
            ),
            (
                "Physician Medical Documents",
                "Verify document type filters and Go Home Slip, Prescription, NFG issuance",
                "1. Filter by document type and status.\n2. Issue Go Home Slip, Prescription, and NFG Medical Clearance Form.",
                "Physician account",
            ),
            (
                "Physician Reports",
                "Verify physician reports and exports",
                "1. Open Reports.\n2. Apply filters.\n3. Confirm daily/monthly consultation, medical certificate, and patient consultation history reports.\n4. Use Print/PDF and CSV export.",
                "Physician account",
            ),
            (
                "Physician Announcements",
                "Verify physician announcements feed is read-only",
                "1. Open Announcements.\n2. Review published feed.\n3. Confirm Add/Edit/Publish controls are absent.",
                "Physician account",
            ),
            (
                "Physician Duty Management",
                "Verify physician Start Duty and End Duty workflow",
                "1. Test Start Duty? and End Duty? confirmations.\n2. Confirm Available and Not Available badge changes and duty toasts.",
                "Physician account",
            ),
            (
                "Physician Break",
                "Verify physician break overlay and Resume Work",
                "1. Go on break and confirm overlay blocks interaction.\n2. Resume Work with confirmation.",
                "Physician account",
            ),
            (
                "Physician Session",
                "Verify physician idle session warning and lock",
                "1. Remain inactive for 4 minutes.\n2. Observe Session Inactivity Detected and session lock behavior.",
                "Physician account",
            ),
            (
                "Physician Settings",
                "Verify physician profile, notifications, and staff schedule",
                "1. Open Profile and Settings.\n2. Update weekly availability and save.\n3. Observe Staff Schedule Updated toast.",
                "Physician account",
            ),
            (
                "Physician Negative",
                "Verify completed consultation cannot be edited",
                "1. Open completed consultation.\n2. Attempt to change clinical fields.",
                "Completed consultation",
            ),
            (
                "Physician Negative",
                "Verify physician cannot access Clinic Staff page",
                "1. Open /admin/user-management/staff manually.\n2. Observe redirect.",
                "Physician account",
            ),
            (
                "Physician Responsive",
                "Verify physician mobile sidebar and consultation page usability",
                "1. Resize below 768px.\n2. Open sidebar and live consultation page.\n3. Confirm layout remains usable.",
                "Mobile viewport; Physician account",
            ),
        ],
    ),
    (
        "DENTIST STAFF",
        [
            (
                "Dentist Navigation",
                "Verify dentist sidebar shows clinical and queue items",
                "1. Sign in as Dentist.\n2. Confirm Dashboard, Reports, Announcements, Queue Management, Patient Records, Consultations, and Medical Documents appear.\n3. Confirm Consultation Requests is absent.",
                "Dentist account",
            ),
            (
                "Dentist Dashboard",
                "Verify dentist dashboard loads",
                "1. Sign in as Dentist.\n2. Open Dashboard.",
                "Dentist account",
            ),
            (
                "Dentist Dashboard",
                "Verify Open queue navigates to dental queue",
                "1. Click Open queue.\n2. Confirm title Dental queue.",
                "Dentist account",
            ),
            (
                "Dentist Dashboard",
                "Verify KPI cards for waiting, ongoing, completed, appointments, schedule",
                "1. Review dentist dashboard KPI cards.",
                "Dentist account",
            ),
            (
                "Dentist Queue Management",
                "Verify dental queue summary cards and search",
                "1. Open Queue Management.\n2. Review summary cards.\n3. Search using dentist patient search.",
                "Dentist account",
            ),
            (
                "Dentist Queue Management",
                "Verify Register walk-in is hidden for dentist",
                "1. Open Queue Management as Dentist.\n2. Confirm Register walk-in is not shown.",
                "Dentist account",
            ),
            (
                "Dentist Queue Management",
                "Verify Start dental consultation with confirmation",
                "1. Open Queue actions.\n2. Click Start consultation.\n3. Confirm Start Consultation?.\n4. Observe navigation to /dentist/consultation/[appointmentId].",
                "Waiting patient in dental queue",
            ),
            (
                "Dentist Queue Management",
                "Verify Call, Complete, Skip, and Mark no-show actions",
                "1. Test Call / recall, Complete, Skip, and Mark no-show with their confirmation dialogs and toasts.",
                "Dental queue row",
            ),
            (
                "Dentist Public Queue Display",
                "Verify public display updates when dentist calls patient",
                "1. Open /queue-management/display on a second screen.\n2. As Dentist, call a patient from dental queue.\n3. Confirm display board updates.",
                "Dentist account; display open",
            ),
            (
                "Dentist Public Queue Display",
                "Verify public display shows recently served after dentist completes patient",
                "1. Complete a patient from dental queue.\n2. View /queue-management/display.\n3. Confirm patient appears in recently served.",
                "Dentist account; display open",
            ),
            (
                "Dentist Public Queue Display",
                "Verify public display reflects clinic break state if shown",
                "1. Put dentist on break.\n2. View /queue-management/display.\n3. Observe break-related display behavior if applicable.",
                "Dentist on break; display open",
            ),
            (
                "Dentist Consultations",
                "Verify View opens consultation summary dialog",
                "1. Click View on a dental consultation.\n2. Review summary dialog content.",
                "Any listed dental consultation",
            ),
            (
                "Dentist Consultations",
                "Verify dentist consultations page loads",
                "1. Open Consultations.\n2. Review stat cards and consultation list.",
                "Dentist account",
            ),
            (
                "Dentist Consultations",
                "Verify dentist live consultation page opens",
                "1. Open a dental consultation from list or queue.\n2. Observe /dentist/consultation/[appointmentId].",
                "Dentist account",
            ),
            (
                "Dentist Consultations",
                "Verify dentist can save and complete dental consultation",
                "1. Enter dental consultation information.\n2. Save draft.\n3. Complete consultation with confirmation.\n4. Observe Consultation Saved and Consultation Completed toasts.",
                "Ongoing dental consultation",
            ),
            (
                "Dentist Consultations",
                "Verify completed dental consultation is read-only",
                "1. Re-open a completed dental consultation.\n2. Confirm fields are read-only.",
                "Completed dental consultation",
            ),
            (
                "Dentist Patient Records",
                "Verify dentist patient records view access",
                "1. Open Patient Records.\n2. Use Profile, History, and Documents actions.",
                "Dentist account",
            ),
            (
                "Dentist Patient Records",
                "Verify Update medical is hidden for dentist",
                "1. Open Patient Records as Dentist.\n2. Confirm Update medical action is not shown.",
                "Dentist account",
            ),
            (
                "Dentist Patient Records",
                "Verify invalid patient search shows empty state",
                "1. Search invalid ID on Patient Records.\n2. Observe No patients found.",
                "Search: 0000-000000",
            ),
            (
                "Dentist Medical Documents",
                "Verify dentist medical documents access",
                "1. Open Medical Documents.\n2. Search, view, preview, print, and issue dental-related documents as available.",
                "Dentist account",
            ),
            (
                "Dentist Reports",
                "Verify dentist reports with dental-focused filters",
                "1. Open Reports.\n2. Confirm consultation type is dental-focused.\n3. Review daily/monthly dental, dental certificate, and patient dental history reports.\n4. Use Print/PDF and CSV export.",
                "Dentist account",
            ),
            (
                "Dentist Announcements",
                "Verify dentist announcements feed is read-only",
                "1. Open Announcements.\n2. Review published feed only.\n3. Confirm management controls are absent.",
                "Dentist account",
            ),
            (
                "Dentist Duty Management",
                "Verify dentist Start Duty and End Duty workflow",
                "1. Test Start Duty? and End Duty? confirmations.\n2. Confirm badge and toast updates.",
                "Dentist account",
            ),
            (
                "Dentist Break",
                "Verify dentist break overlay and Resume Work",
                "1. Go on break and confirm overlay blocks interaction.\n2. Resume Work with confirmation.",
                "Dentist account",
            ),
            (
                "Dentist Session",
                "Verify dentist idle session warning and lock",
                "1. Remain inactive for 4 minutes.\n2. Observe Session Inactivity Detected and session lock behavior.",
                "Dentist account",
            ),
            (
                "Dentist Settings",
                "Verify dentist profile, notifications, and staff schedule",
                "1. Open Profile and Settings.\n2. Update weekly availability and save.",
                "Dentist account",
            ),
            (
                "Dentist Negative",
                "Verify dentist cannot access admin staff management",
                "1. Open /admin/user-management/staff manually.\n2. Observe redirect.",
                "Dentist account",
            ),
            (
                "Dentist Negative",
                "Verify queue skip cancel leaves patient in queue",
                "1. Open Skip Patient? dialog.\n2. Click Cancel.\n3. Confirm patient remains in previous status.",
                "Waiting patient in dental queue",
            ),
            (
                "Dentist Responsive",
                "Verify dentist mobile sidebar and dental queue usability",
                "1. Resize below 768px.\n2. Open sidebar and dental queue.\n3. Confirm controls remain accessible.",
                "Mobile viewport; Dentist account",
            ),
        ],
    ),
]


def resolve_template() -> Path:
    for candidate in TEMPLATE_CANDIDATES:
        if candidate.exists():
            return candidate
    raise FileNotFoundError("Could not find CampusCare test case template DOCX.")


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)


def add_section_row(table, section: str) -> None:
    row = table.add_row()
    cells = row.cells
    set_cell_text(cells[0], "")
    set_cell_text(cells[1], section)
    set_cell_text(cells[2], section)
    set_cell_text(cells[3], "")
    set_cell_text(cells[4], "")
    set_cell_text(cells[5], "")
    set_cell_text(cells[6], "")
    set_cell_text(cells[7], "")


def add_test_row(
    table, tc_id: str, module: str, scenario: str, action: str, actual_input: str
) -> None:
    row = table.add_row()
    cells = row.cells
    set_cell_text(cells[0], tc_id)
    set_cell_text(cells[1], module)
    set_cell_text(cells[2], scenario)
    set_cell_text(cells[3], action)
    set_cell_text(cells[4], actual_input)
    set_cell_text(cells[5], "")
    set_cell_text(cells[6], "")
    set_cell_text(cells[7], "")


def update_metadata_table(doc: Document, total_cases: int) -> None:
    meta = doc.tables[0]
    today = date.today().strftime("%B %d, %Y")

    for cell in meta.rows[2].cells:
        if cell.text.strip() and "Written Date" not in cell.text:
            cell.text = today

    label_found = False
    for cell in meta.rows[6].cells:
        if "Total" in cell.text:
            label_found = True
            continue
        if label_found and not cell.text.strip().startswith("Rating"):
            cell.text = str(total_cases)
            break

    # Update first section label in metadata table from GETTING STARTED to LANDING PAGE
    if len(meta.rows) > 10:
        for cell in meta.rows[10].cells:
            if "GETTING STARTED" in cell.text:
                cell.text = cell.text.replace("GETTING STARTED", "LANDING PAGE")


def clear_table_rows(table) -> None:
    while len(table.rows) > 0:
        table._tbl.remove(table.rows[0]._tr)


def main() -> None:
    template = resolve_template()
    doc = Document(str(template))
    test_table = doc.tables[1]
    clear_table_rows(test_table)

    tc_num = 0
    for section, cases in SECTIONS:
        add_section_row(test_table, section)
        for module, scenario, action, actual_input in cases:
            tc_num += 1
            add_test_row(
                test_table,
                f"TC-{tc_num:03d}",
                module,
                scenario,
                action,
                actual_input,
            )

    update_metadata_table(doc, tc_num)

    OUTPUT_PROJECT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PROJECT))
    doc.save(str(OUTPUT_DOWNLOADS))
    print(f"Generated {tc_num} test cases across {len(SECTIONS)} role sections")
    print(f"Template: {template}")
    print(f"Saved: {OUTPUT_PROJECT}")
    print(f"Saved: {OUTPUT_DOWNLOADS}")


if __name__ == "__main__":
    main()
